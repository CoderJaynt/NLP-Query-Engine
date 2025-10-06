from flask import Flask, request, jsonify, render_template
from sqlalchemy import create_engine, inspect, text
import os
import asyncio
from PyPDF2 import PdfReader
from docx import Document
from flask_caching import Cache
import hashlib, json, time
from groq import Groq



app = Flask(__name__)
# Configure simple in-memory cache
app.config["CACHE_TYPE"] = "SimpleCache"
app.config["CACHE_DEFAULT_TIMEOUT"] = 60  # 1 minute default
cache = Cache(app)

# Initialize Groq client
groq_client = Groq(api_key="") #### ****Give your Groq API Key(FREE)****

engines = {}
def get_engine(conn_str):
    """Create or reuse an engine with a connection pool."""
    if conn_str not in engines:
        engines[conn_str] = create_engine(
            conn_str,
            pool_size=5,          # max persistent connections
            max_overflow=2,       # extra temporary connections
            pool_recycle=1800,    # recycle every 30 mins
            pool_pre_ping=True    # check connection before use
        )
    return engines[conn_str]

# ---- Step 3: batch processing helper ----
def batch_process_embeddings(texts, batch_size=8):
    """
    Dummy batch processor — simulates embedding generation
    so the code runs without an actual model.
    """
    embeddings = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        # Simulate "embedding" as text length
        emb_batch = [len(t) for t in batch]
        embeddings.extend(emb_batch)
    return embeddings

def interpret_user_query(query, schema_desc):
    """
    Send user query and schema to ChatGroq to get SQL query or document query intent.
    Supports both structured (DB) and unstructured (PDF/DOCX/TXT) data.
    """
    prompt = f"""
    You are an AI Query Assistant that handles both structured and unstructured data.

     For database-related queries:
       - Use the provided schema to generate valid MySQL SQL queries.
       - Follow clause order: SELECT → FROM → JOIN → WHERE → GROUP BY → ORDER BY → LIMIT.
       - Only use tables/columns present in the schema.
       - Return your result in strict JSON format:
         {{
            "type": "sql",
            "query": "SELECT ...",
            "explanation": "..."
         }}

     For document or unstructured data queries:
       - Understand the user's intent based on the question.
       - Possible intents include:
         • summarize → summarize the document
         • search → find mentions of specific terms
         • keyword → extract relevant topics or entities
       - Return the result in this JSON format:
         {{
            "type": "document",
            "intent": "summarize" / "search" / "keyword",
            "keywords": ["..."]
         }}

    Database Schema (for reference):
    {schema_desc}

    User Question:
    {query}

    Important:
    - Always reply with valid JSON only.
    - Do NOT include markdown, ```json blocks, or natural language outside JSON.
    """

    response = groq_client.chat.completions.create(
        model="gemma2-9b-it",  # or mixtral-8x7b-32768
        messages=[{"role": "user", "content": prompt}]
    )

    content = response.choices[0].message.content.strip()

    # --- Clean and parse safely ---
    try:
        if "```" in content:
            content = content.split("```")[1]
            content = content.replace("json", "").strip()
        return json.loads(content)
    except json.JSONDecodeError:
        # fallback if Groq returns text
        return {"type": "error", "raw_output": content}





@app.route("/api/async-query", methods=["POST"])
async def async_query():
    data = await request.get_json()
    conn_str = data.get("connection_string")
    query = data.get("query")
    engine = get_engine(conn_str)

    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, lambda: engine.execute(text(query)).fetchall())
    return jsonify([dict(r._mapping) for r in result])

@app.route("/api/query/paged", methods=["POST"])
def query_paged():
    data = request.get_json()
    connection_string = data.get("connection_string")
    query_text = data.get("query", "").lower()
    limit = int(data.get("limit", 20))
    offset = int(data.get("offset", 0))

    engine = get_engine(connection_string)
    conn = engine.connect()

    if "employee" in query_text:
        sql_query = f"SELECT * FROM employees LIMIT {limit} OFFSET {offset};"
    else:
        sql_query = f"SELECT * FROM departments LIMIT {limit} OFFSET {offset};"

    result = conn.execute(text(sql_query))
    rows = [dict(row._mapping) for row in result]
    conn.close()

    return jsonify({
        "query": sql_query,
        "limit": limit,
        "offset": offset,
        "results": rows
    })


@app.route("/")
def home():
    return render_template("index.html")

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

uploaded_files = []

@app.route("/api/ingest/documents", methods=["POST"])
def ingest_documents():
    try:
        global uploaded_files
        uploaded_files.clear()  # clear previous batch
        if "files" not in request.files:
            return jsonify({"error": "No files provided"}), 400

        files = request.files.getlist("files")
        extracted_data = []

        for file in files:
            filename = file.filename
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(file_path)
            uploaded_files.append(filename)

            text_content = ""
            if filename.lower().endswith(".pdf"):
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text_content += page.extract_text() or ""
            elif filename.lower().endswith(".docx"):
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
            elif filename.lower().endswith(".txt") or filename.lower().endswith(".csv"):
                text_content = file.read().decode("utf-8")

            extracted_data.append({
                "filename": filename,
                "text_preview": text_content[:300] + "..." if len(text_content) > 300 else text_content
            })

        return jsonify({
            "message": f"{len(files)} file(s) uploaded successfully.",
            "uploaded_files": uploaded_files
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500



# --------------------------
# Existing schema discovery endpoint
# --------------------------
@app.route("/api/ingest/database", methods=["POST"])
def ingest_database():
    data = request.get_json()
    connection_string = data.get("connection_string")
    try:
        engine = get_engine(connection_string)
        inspector = inspect(engine)
        schema_info = []

        for table_name in inspector.get_table_names():
            # Extract columns
            columns = [
                {"name": col["name"], "type": str(col["type"])}
                for col in inspector.get_columns(table_name)
            ]
        # Extract foreign keys
        raw_fks = inspector.get_foreign_keys(table_name)
        formatted_fks = []
        for fk in raw_fks:
            formatted_fks.append({
                "column": fk.get("constrained_columns", []),
                "ref_table": fk.get("referred_table", ""),
                "ref_column": fk.get("referred_columns", [])
            })
        schema_info.append({
            "name": table_name,
            "columns": columns,
            "foreign_keys": formatted_fks
        })

        return jsonify({"tables": schema_info})

    except Exception as e:
            return jsonify({"error": str(e)}), 500



# --------------------------
# New Query Endpoint
# --------------------------
@app.route("/api/query", methods=["POST"])
def query_data():
    data = request.get_json()
    connection_string = data.get("connection_string", "")
    user_query = data.get("query", "").strip().lower()

    # --- Cache key
    cache_key = hashlib.md5((connection_string + user_query).encode("utf-8")).hexdigest()

    cached = cache.get(cache_key)
    if cached:
        cached["cache_hit"] = True
        return jsonify(cached)

    try:
        # --- Get schema info if DB connection exists
        schema = {}
        engine = None
        if connection_string:
            engine = get_engine(connection_string)
            inspector = inspect(engine)
            schema = {t: [c["name"] for c in inspector.get_columns(t)] for t in inspector.get_table_names()}

        # --- Step 1: Ask Groq what to do
        groq_result = interpret_user_query(user_query, json.dumps(schema))

        # --- Case 1: SQL Query ---
        if groq_result["type"] == "sql":
            conn = engine.connect()
            result = conn.execute(text(groq_result["query"]))
            rows = [dict(r._mapping) for r in result]
            conn.close()

            response = {
                "mode": "database",
                "sql": groq_result["query"],
                "results": rows,
                "explanation": groq_result.get("explanation", ""),
                "cache_hit": False
            }
            cache.set(cache_key, response)
            return jsonify(response)

        # --- Case 2: Document Query ---
        elif groq_result.get("type") == "document":
            # --- Safe parsing from Groq ---
            intent = groq_result.get("intent", "")
            keywords = groq_result.get("keywords", [])
            matched_docs = []

            # --- Intent fallback ---
            if not intent:
                q = user_query.lower()
                if any(x in q for x in ["summarize", "summary"]):
                    intent = "summarize"
                elif any(x in q for x in ["find", "search", "where", "contains", "show", "look for"]):
                    intent = "search"
                else:
                    intent = "qa"  # general question answering fallback

            for filename in uploaded_files:
                file_path = os.path.join(UPLOAD_FOLDER, filename)
                text_content = ""

                # --- Extract text from different file types ---
                if filename.lower().endswith(".pdf"):
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        text_content += page.extract_text() or ""
                elif filename.lower().endswith(".docx"):
                    doc = Document(file_path)
                    for para in doc.paragraphs:
                        text_content += para.text + "\n"
                elif filename.lower().endswith(".txt"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        text_content = f.read()

                if not text_content.strip():
                    continue  # skip empty docs

                # ---- Intent 1: Summarize the document ----
                if intent == "summarize":
                    try:
                        summary_prompt = f"Summarize this document briefly:\n{text_content[:3000]}"
                        summary_response = groq_client.chat.completions.create(
                            model="gemma2-9b-it",
                            messages=[{"role": "user", "content": summary_prompt}]
                        )
                        summary = getattr(summary_response.choices[0].message, "content", "").strip()
                        matched_docs.append({"filename": filename, "summary": summary})
                    except Exception as e:
                        matched_docs.append({"filename": filename, "error": f"Summarization failed: {str(e)}"})

                # ---- Intent 2: Search by keywords ----
                elif intent == "search" and keywords:
                    for kw in keywords:
                        if kw.lower() in text_content.lower():
                            idx = text_content.lower().find(kw.lower())
                            snippet = text_content[max(0, idx - 50):idx + 50].replace("\n", " ")
                            matched_docs.append({
                                "filename": filename,
                                "keyword": kw,
                                "snippet": snippet
                            })

                # ---- Intent 3: Question Answering / Generic Query ----
                elif intent == "qa":
                    try:
                        doc_prompt = f"Answer the question '{user_query}' using the content below:\n{text_content[:8000]}"
                        doc_response = groq_client.chat.completions.create(
                            model="gemma2-9b-it",
                            messages=[{"role": "user", "content": doc_prompt}]
                        )
                        answer = getattr(doc_response.choices[0].message, "content", "").strip()
                        matched_docs.append({"filename": filename, "answer": answer})
                    except Exception as e:
                        matched_docs.append({"filename": filename, "error": f"QA failed: {str(e)}"})

            # --- Final Response for Documents ---
            final_response = {
                "mode": "document",
                "query": user_query,
                "intent": intent,
                "results": matched_docs,
                "cache_hit": False
            }

            cache.set(cache_key, final_response)
            return jsonify(final_response)

        else:
            return jsonify({"error": "Groq could not interpret query type"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500


    
@app.route("/api/query/documents", methods=["POST"])
def query_documents():
    data = request.get_json()
    query_text = data.get("query", "").lower()

    if not query_text:
        return jsonify({"error": "Query text is required"}), 400

    try:
        results = []

        # Iterate over all files in uploads folder
        for filename in os.listdir(UPLOAD_FOLDER):
            file_path = os.path.join(UPLOAD_FOLDER, filename)

            # Read the file text
            text_content = ""
            if filename.lower().endswith(".pdf"):
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text_content += page.extract_text() or ""
            elif filename.lower().endswith(".docx"):
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
            elif filename.lower().endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()

            # Simple keyword search
            if query_text in text_content.lower():
                # Return a preview snippet of 100 chars around first match
                idx = text_content.lower().find(query_text)
                start = max(0, idx - 50)
                end = min(len(text_content), idx + 50)
                snippet = text_content[start:end].replace("\n", " ")
                results.append({
                    "filename": filename,
                    "snippet": snippet
                })

        return jsonify({
            "query": query_text,
            "matches": results
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500



if __name__ == "__main__":
    app.run(debug=True)